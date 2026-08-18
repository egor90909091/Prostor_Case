"""
Проверка структуры .docx, который собирает Prostor.Tz/DocxWriter.cs.

Питон здесь повторяет ровно ту же разметку WordprocessingML, что и C#.
Тест доказывает, что такой пакет корректно читается сторонним парсером
(python-docx) — то есть Word/LibreOffice его тоже откроют.
"""
import io, zipfile, sys
from xml.sax.saxutils import escape

MAIN = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL  = "http://schemas.openxmlformats.org/package/2006/relationships"
DOCR = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

CONTENT_TYPES = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
<Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>'''

ROOT_RELS = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="{REL}">
<Relationship Id="rId1" Type="{DOCR}/officeDocument" Target="word/document.xml"/>
</Relationships>'''

DOC_RELS = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="{REL}">
<Relationship Id="rId1" Type="{DOCR}/styles" Target="styles.xml"/>
</Relationships>'''

STYLES = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="{MAIN}">
<w:docDefaults><w:rPrDefault><w:rPr>
  <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>
  <w:sz w:val="24"/><w:szCs w:val="24"/><w:lang w:val="ru-RU"/>
</w:rPr></w:rPrDefault></w:docDefaults>
<w:style w:type="paragraph" w:default="1" w:styleId="Normal">
  <w:name w:val="Normal"/><w:qFormat/>
</w:style>
<w:style w:type="paragraph" w:styleId="Heading1">
  <w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:qFormat/>
  <w:pPr><w:spacing w:before="240" w:after="120"/><w:outlineLvl w:val="0"/></w:pPr>
  <w:rPr><w:b/><w:sz w:val="32"/></w:rPr>
</w:style>
<w:style w:type="paragraph" w:styleId="Heading2">
  <w:name w:val="heading 2"/><w:basedOn w:val="Normal"/><w:qFormat/>
  <w:pPr><w:spacing w:before="200" w:after="100"/><w:outlineLvl w:val="1"/></w:pPr>
  <w:rPr><w:b/><w:sz w:val="26"/></w:rPr>
</w:style>
</w:styles>'''


def heading(text, level, center=False):
    jc = '<w:jc w:val="center"/>' if center else ''
    return (f'<w:p><w:pPr><w:pStyle w:val="Heading{level}"/>{jc}</w:pPr>'
            f'<w:r><w:t xml:space="preserve">{escape(text)}</w:t></w:r></w:p>')


def paragraph(text, center=False, bold=False, italic=False):
    jc = '<w:jc w:val="center"/>' if center else '<w:jc w:val="both"/>'
    rpr = ('<w:b/>' if bold else '') + ('<w:i/>' if italic else '')
    rpr = f'<w:rPr>{rpr}</w:rPr>' if rpr else ''
    return (f'<w:p><w:pPr>{jc}<w:spacing w:after="120"/></w:pPr>'
            f'<w:r>{rpr}<w:t xml:space="preserve">{escape(text)}</w:t></w:r></w:p>')


SECTIONS = [
    ("Цели и задачи работ", "Выполнение работ по услуге «Проектно-технический документ»."),
    ("Периметр выполнения работ", "Объект работ: Восточно-Мыгинское месторождение."),
    ("Сроки выполнения работ", "Начало работ: 01.09.2026. Окончание работ: 30.11.2026. Общая продолжительность: 91 календарных дней."),
    ("Содержание работ, особенности их выполнения и результаты",
     "Этап 1. ПТД Восточно-Мыгинского месторождения Продолжительность по типовому расчёту: 138 дн.\n"
     "Этап 2. Оперативный подсчёт геологических запасов. Отчётные материалы: Информационный отчет."),
    ("Требования к документации", "Результаты работ передаются информационными отчётами: DOC/DOCX и PDF."),
    ("Условия привлечения субподрядчиков", "Исполнители работ: NNG, MRNG."),
]


def build():
    body = []
    body.append(heading("ТЕХНИЧЕСКОЕ ЗАДАНИЕ", 1, center=True))
    body.append(paragraph("ТЗ на проектно-технический документ (ПТД)", center=True, bold=True))
    body.append(paragraph("Услуга: Проектно-технический документ", center=True))
    body.append(paragraph("Объект работ: Восточно-Мыгинское месторождение", center=True))
    body.append(paragraph("Заказчик: {Полное-Наименование-ДО-Заказчика}"))
    body.append(paragraph(""))
    for i, (title, text) in enumerate(SECTIONS, start=1):
        body.append(heading(f"{i}. {title}", 2))
        for line in [x for x in text.split("\n") if x.strip()]:
            body.append(paragraph(line.strip()))
    body.append(heading("Приложение. Оценка качества технического задания", 2))
    body.append(paragraph("Готовность к согласованию: 78%.", bold=True))
    body.append(paragraph("Перед согласованием необходимо устранить: не указан объект работ."))
    body.append(paragraph("— [внимание] Заявленный срок ниже типового. Сравните срок с историей."))
    body.append(paragraph(""))
    body.append(paragraph("ПОДПИСИ СТОРОН:", bold=True))
    body.append(paragraph("Заказчик: ______________________        Исполнитель: ______________________"))
    body.append('<w:sectPr><w:pgSz w:w="11906" w:h="16838"/>'
                '<w:pgMar w:top="1134" w:right="850" w:bottom="1134" w:left="1701" '
                'w:header="708" w:footer="708" w:gutter="0"/></w:sectPr>')

    document = ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                f'<w:document xmlns:w="{MAIN}"><w:body>{"".join(body)}</w:body></w:document>')

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", CONTENT_TYPES)
        z.writestr("_rels/.rels", ROOT_RELS)
        z.writestr("word/_rels/document.xml.rels", DOC_RELS)
        z.writestr("word/styles.xml", STYLES)
        z.writestr("word/document.xml", document)
    return buffer.getvalue()


def main():
    data = build()
    open("/tmp/tz_sample.docx", "wb").write(data)

    import docx
    d = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in d.paragraphs if p.text.strip()]
    styles = {p.style.name for p in d.paragraphs}

    assert paragraphs[0] == "ТЕХНИЧЕСКОЕ ЗАДАНИЕ", paragraphs[:3]
    assert any("Готовность к согласованию: 78%" in p for p in paragraphs)
    assert any(s.startswith("Heading") or s == "heading 1" or s == "heading 2" for s in styles), styles
    assert len(paragraphs) >= 15, len(paragraphs)

    from xml.etree import ElementTree
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        for name in ("[Content_Types].xml", "_rels/.rels",
                     "word/_rels/document.xml.rels", "word/styles.xml", "word/document.xml"):
            ElementTree.fromstring(z.read(name))     # XML валиден

    print(f"OK: {len(data)} байт, {len(paragraphs)} абзацев, стили {sorted(styles)}")
    print("первые абзацы:", paragraphs[:4])


if __name__ == "__main__":
    main()
